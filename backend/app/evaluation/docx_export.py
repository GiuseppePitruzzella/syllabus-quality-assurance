"""Institutional DOCX export for persisted syllabus evaluations."""
from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
import json
import re
from typing import Any, Iterable
import unicodedata

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.evaluation.synthesizer import (
    CRITERION_NAMES,
    CRITERION_RECOMMENDATIONS,
)


UNICT_BLUE = RGBColor(0x00, 0x45, 0x7C)
UNICT_LIGHT_BLUE = "EAF2F8"
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x5B, 0x67, 0x75)
CRITICAL = RGBColor(0xA6, 0x1B, 0x2B)
IMPROVABLE = RGBColor(0x8A, 0x5A, 0x00)
ADEQUATE = RGBColor(0x1D, 0x6B, 0x46)

CORE_ORDER = tuple(f"C{i}" for i in range(1, 10))
EXTENDED_ORDER = tuple(f"E{i}" for i in range(1, 6))

EXTENDED_NAMES = {
    "E1": "Allineamento con SUA-CdS",
    "E2": "Allineamento con Matrice di Tuning",
    "E3": "Coerenza con Regolamento didattico",
    "E4": "Coerenza cross-lingua",
    "E5": "Aderenza agli usi dipartimentali / di CdL",
}

SECTION_FALLBACK_CRITERIA: dict[str, tuple[str, ...]] = {
    "learning_outcomes": ("C2", "C3", "C4"),
    "teaching_methods": ("C1",),
    "prerequisites": ("C5",),
    "attendance": ("C1",),
    "course_content": ("C6", "C7"),
    "references": ("C9",),
    "schedule": ("C7",),
    "assessment_methods": ("C6", "C8"),
    "sample_questions": ("C6", "C8"),
}


@dataclass(frozen=True)
class ExternalDocumentExport:
    criterion_code: str
    title: str | None
    document_type: str
    document_version: int
    resolution_reason: str


@dataclass(frozen=True)
class EvaluationExportBundle:
    evaluation: Any
    syllabus: Any
    cdl: Any
    department: Any
    external_documents: tuple[ExternalDocumentExport, ...] = ()


@dataclass(frozen=True)
class Annotation:
    code: str
    score: int
    justification: str
    evidence_text: str | None
    source_field: str | None


def build_evaluation_docx(bundle: EvaluationExportBundle) -> bytes:
    """Build one polished, non-technical DOCX for a persisted evaluation."""
    doc = Document()
    _configure_document(doc)
    doc.core_properties.title = (
        f"Valutazione del syllabus - {bundle.evaluation.course_name_snapshot}"
    )
    doc.core_properties.author = "Syllabus Quality Assurance"
    doc.core_properties.subject = "Valutazione automatica della qualità del syllabus"
    _add_cover(doc, bundle)
    _add_summary(doc, bundle)
    _add_core_results(doc, bundle)
    _add_extended_results(doc, bundle)
    _add_external_documents(doc, bundle.external_documents)
    _add_method_note(doc)
    _add_annotated_syllabus(doc, bundle)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_filename(bundle: EvaluationExportBundle) -> str:
    code = _filename_part(str(getattr(bundle.cdl, "code", None) or "CdS").upper())
    course = _filename_part(bundle.evaluation.course_name_snapshot)
    stamp = _as_date(bundle.evaluation.started_at).isoformat()
    return f"{code}__{course}__{stamp}__evaluation.docx"


def cdl_zip_filename(cdl: Any) -> str:
    code = _filename_part(str(cdl.code or "CdS").upper())
    return f"{code}__valutazioni__{date.today().isoformat()}.zip"


class EvaluationExportNotFoundError(LookupError):
    """Raised when an export target has no eligible persisted evaluation."""


def load_evaluation_export_bundle(session: Any, evaluation_uuid: str) -> EvaluationExportBundle:
    """Load one exact terminal, non-failed evaluation for export."""
    from app.models import (
        CorsoDiLaurea,
        Department,
        EvaluationExternalDocument,
        EvaluationResult,
        LocalDocument,
        Syllabus,
    )

    row = (
        session.query(EvaluationResult, Syllabus, CorsoDiLaurea, Department)
        .join(Syllabus, EvaluationResult.syllabus_id == Syllabus.id)
        .join(CorsoDiLaurea, Syllabus.cdl_id == CorsoDiLaurea.id)
        .join(Department, CorsoDiLaurea.department_id == Department.id)
        .filter(
            EvaluationResult.evaluation_uuid == evaluation_uuid,
            EvaluationResult.status.in_(("completed", "partial")),
        )
        .one_or_none()
    )
    if row is None:
        raise EvaluationExportNotFoundError(
            f"evaluation not exportable: evaluation_uuid={evaluation_uuid!r}"
        )
    evaluation, syllabus, cdl, department = row
    external_rows = (
        session.query(EvaluationExternalDocument, LocalDocument)
        .outerjoin(
            LocalDocument,
            EvaluationExternalDocument.local_document_id == LocalDocument.id,
        )
        .filter(EvaluationExternalDocument.evaluation_result_id == evaluation.id)
        .order_by(
            EvaluationExternalDocument.criterion_code,
            EvaluationExternalDocument.local_document_id,
        )
        .all()
    )
    return EvaluationExportBundle(
        evaluation=evaluation,
        syllabus=syllabus,
        cdl=cdl,
        department=department,
        external_documents=tuple(
            ExternalDocumentExport(
                criterion_code=audit.criterion_code,
                title=document.title if document is not None else None,
                document_type=audit.document_type_snapshot,
                document_version=audit.document_version_snapshot,
                resolution_reason=audit.resolution_reason,
            )
            for audit, document in external_rows
        ),
    )


def load_cdl_export_bundles(
    session: Any,
    cdl_id: int,
) -> tuple[Any, list[EvaluationExportBundle]]:
    """Load the latest terminal, non-failed run for each syllabus in a CdS."""
    from app.models import CorsoDiLaurea, EvaluationResult, Syllabus

    cdl = session.query(CorsoDiLaurea).filter(CorsoDiLaurea.id == cdl_id).one_or_none()
    if cdl is None:
        raise EvaluationExportNotFoundError(f"CdS not found: cdl_id={cdl_id}")
    rows = (
        session.query(EvaluationResult)
        .join(Syllabus, EvaluationResult.syllabus_id == Syllabus.id)
        .filter(
            Syllabus.cdl_id == cdl_id,
            EvaluationResult.status.in_(("completed", "partial")),
        )
        .order_by(
            EvaluationResult.syllabus_id,
            EvaluationResult.started_at.desc(),
            EvaluationResult.id.desc(),
        )
        .all()
    )
    latest: dict[int, EvaluationResult] = {}
    for evaluation in rows:
        latest.setdefault(evaluation.syllabus_id, evaluation)
    bundles = [
        load_evaluation_export_bundle(session, evaluation.evaluation_uuid)
        for evaluation in latest.values()
    ]
    return cdl, bundles


def _configure_document(doc: DocxDocument) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Title": (26, UNICT_BLUE, 0, 10),
        "Heading 1": (17, UNICT_BLUE, 16, 7),
        "Heading 2": (13, UNICT_BLUE, 12, 5),
        "Heading 3": (11, INK, 8, 3),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("UNIVERSITÀ DEGLI STUDI DI CATANIA · SYLLABUS QUALITY ASSURANCE")
    _format_run(run, size=8, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Documento generato automaticamente · Supporto alla revisione")
    _format_run(run, size=8, color=MUTED)


def _add_cover(doc: DocxDocument, bundle: EvaluationExportBundle) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("UNIVERSITÀ DEGLI STUDI DI CATANIA")
    _format_run(run, size=11, color=UNICT_BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Valutazione della qualità del syllabus")
    _format_run(run, size=26, color=UNICT_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(bundle.evaluation.course_name_snapshot)
    _format_run(run, size=17, color=INK, bold=True)

    metadata = [
        ("Corso di Studio", f"{str(bundle.cdl.code).upper()} · {bundle.cdl.name}"),
        ("Dipartimento", bundle.department.name),
        ("Docente", bundle.syllabus.teacher or "—"),
        (
            "Anno accademico",
            bundle.syllabus.academic_year
            or getattr(bundle.cdl, "academic_year", None)
            or "—",
        ),
        ("Stato valutazione", _status_label(bundle.evaluation.status)),
        ("Data valutazione", _format_date(bundle.evaluation.started_at)),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for label, value in metadata:
        cells = table.add_row().cells
        cells[0].width = Inches(1.65)
        cells[1].width = Inches(4.95)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_run = cells[0].paragraphs[0].add_run(label)
        _format_run(label_run, size=9, color=MUTED, bold=True)
        value_run = cells[1].paragraphs[0].add_run(str(value))
        _format_run(value_run, size=10, color=INK)
        _shade_cell(cells[0], UNICT_LIGHT_BLUE)
    _set_table_widths(table, (1.65, 4.95))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Il documento riporta una valutazione automatica di supporto alla revisione. "
        "I criteri estesi E1-E5 sono presentati separatamente e non concorrono al CoreScore."
    )
    _format_run(run, size=9, color=MUTED, italic=True)


def _add_summary(doc: DocxDocument, bundle: EvaluationExportBundle) -> None:
    evaluation = bundle.evaluation
    doc.add_heading("Sintesi della valutazione", level=1)
    values = [
        ("CoreScore", _format_score(evaluation.core_score)),
        ("Copertura", _format_percent(evaluation.coverage)),
        ("Criteri valutati", _evaluated_count(evaluation.criterion_scores)),
    ]
    metrics = doc.add_paragraph()
    metrics.paragraph_format.space_after = Pt(10)
    for index, (label, value) in enumerate(values):
        if index:
            separator = metrics.add_run("    ·    ")
            _format_run(separator, size=11, color=MUTED)
        label_run = metrics.add_run(f"{label} ")
        _format_run(label_run, size=9, color=MUTED, bold=True)
        value_run = metrics.add_run(value)
        _format_run(value_run, size=16, color=UNICT_BLUE, bold=True)

    _add_scorecard(doc, evaluation.criterion_scores)

    priorities = _priority_codes(evaluation.criterion_scores)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    lead = p.add_run("Priorità di revisione: ")
    _format_run(lead, bold=True, color=INK)
    text = ", ".join(
        f"{code} ({'criticità' if score == 0 else 'da migliorare'})"
        for code, score in priorities
    ) or "nessuna criticità o area migliorabile rilevata"
    p.add_run(text)


def _add_scorecard(doc: DocxDocument, criterion_scores: Any) -> None:
    """One scannable C1-C9 row: code, name, score, colour-coded outcome."""
    scores = criterion_scores or {}
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    for idx, label in enumerate(("", "Criterio", "Punteggio", "Esito")):
        cell = table.rows[0].cells[idx]
        cell.text = label
        _shade_cell(cell, UNICT_LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            _format_run(run, size=9, color=UNICT_BLUE, bold=True)
    for code in CORE_ORDER:
        score = _score_value(scores.get(code))
        label, color = _outcome_tokens(score)
        cells = table.add_row().cells
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        code_run = cells[0].paragraphs[0].add_run(code)
        _format_run(code_run, size=9.5, color=UNICT_BLUE, bold=True)
        name_run = cells[1].paragraphs[0].add_run(CRITERION_NAMES.get(code, code))
        _format_run(name_run, size=9.5, color=INK)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_run = cells[2].paragraphs[0].add_run(
            f"{score}/2" if score is not None else "—"
        )
        _format_run(score_run, size=9.5, color=color, bold=True)
        esito_run = cells[3].paragraphs[0].add_run(label)
        _format_run(esito_run, size=9.5, color=color, bold=True)
    _set_table_widths(table, (0.55, 4.0, 0.85, 1.2))


def _add_core_results(doc: DocxDocument, bundle: EvaluationExportBundle) -> None:
    doc.add_heading("Revisione dei criteri C1-C9", level=1)
    judgments = _core_judgments(bundle.evaluation.agent_outputs)
    scores = bundle.evaluation.criterion_scores or {}
    na_reasons = _na_reason_map(bundle.evaluation.na_criteria)
    seen_evidence: set[str] = set()

    # Full detail only for criteria that need attention (0/1) or are NA.
    # Adequate criteria (score 2) are recapped in one line below, since the
    # scorecard already certifies them — no need to restate the obvious.
    needs_review = [
        code for code in CORE_ORDER if _score_value(scores.get(code)) != 2
    ]
    if not needs_review:
        p = doc.add_paragraph()
        _format_run(
            p.add_run(
                "Tutti i criteri valutati risultano adeguati: "
                "nessuna area di revisione segnalata."
            ),
            color=ADEQUATE,
            bold=True,
        )
    for code in needs_review:
        score = _score_value(scores.get(code))
        judgment = judgments.get(code)
        heading = doc.add_paragraph(style="Heading 2")
        heading.add_run(f"{code} · {CRITERION_NAMES.get(code, code)}")
        _add_outcome_run(heading, score)

        if score is None:
            p = doc.add_paragraph()
            p.add_run("Non valutabile. ").bold = True
            p.add_run(_humanize_na_reason(na_reasons.get(code)))
            continue

        if judgment and judgment.get("justification"):
            doc.add_paragraph(_humanize(str(judgment["justification"])))
        evidences = judgment.get("evidences", []) if judgment else []
        rendered = _select_evidences(evidences, seen_evidence)
        if rendered:
            p = doc.add_paragraph()
            p.add_run("Evidenze dal syllabus").bold = True
            for text in rendered:
                quote = doc.add_paragraph(style="Quote")
                quote.add_run(f"“{text}”")
        p = doc.add_paragraph()
        lead = p.add_run("Indicazione per la revisione. ")
        _format_run(lead, bold=True, color=CRITICAL if score == 0 else IMPROVABLE)
        p.add_run(CRITERION_RECOMMENDATIONS.get(code, "Rivedere il criterio."))

    adequate = [code for code in CORE_ORDER if _score_value(scores.get(code)) == 2]
    if adequate:
        doc.add_heading("Criteri adeguati", level=2)
        for code in adequate:
            judgment = judgments.get(code)
            summary = _first_sentence(str(judgment.get("justification") or "")) \
                if judgment else ""
            p = doc.add_paragraph(style="List Bullet")
            lead = p.add_run(f"{code} · {CRITERION_NAMES.get(code, code)}")
            _format_run(lead, bold=True, color=UNICT_BLUE)
            if summary:
                p.add_run(" — ")
                p.add_run(summary)


def _add_extended_results(doc: DocxDocument, bundle: EvaluationExportBundle) -> None:
    raw = bundle.evaluation.extended_criteria_result or {}
    if not raw:
        return
    doc.add_heading("Criteri estesi E1-E5", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Questi criteri non concorrono al CoreScore.")
    _format_run(run, size=9, color=MUTED, italic=True)

    agent_output = raw.get("agent_output") or {}
    judgments = {
        item.get("criterion_code"): item
        for item in agent_output.get("judgments", [])
        if isinstance(item, dict)
    }
    scores = raw.get("criterion_scores") or {}
    na_reasons = {
        item.get("criterion_code"): item.get("reason")
        for item in raw.get("na_criteria", [])
        if isinstance(item, dict)
    }
    for code in EXTENDED_ORDER:
        heading = doc.add_paragraph(style="Heading 2")
        heading.add_run(f"{code} · {EXTENDED_NAMES[code]}")
        score = _score_value(scores.get(code))
        _add_outcome_run(heading, score)
        judgment = judgments.get(code)
        if score is None:
            p = doc.add_paragraph()
            p.add_run("Non valutabile. ").bold = True
            p.add_run(_humanize_na_reason(na_reasons.get(code)))
        elif judgment:
            doc.add_paragraph(_humanize(str(judgment.get("justification") or "—")))
            for evidence in judgment.get("evidences") or []:
                quote = doc.add_paragraph(style="Quote")
                quote.add_run(f"“{_evidence_text(str(evidence.get('text', '')))}”")


def _add_external_documents(
    doc: DocxDocument,
    documents: Iterable[ExternalDocumentExport],
) -> None:
    items = list(documents)
    if not items:
        return
    doc.add_heading("Documenti esterni utilizzati", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    headers = ("Criterio", "Documento", "Versione", "Selezione")
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = label
        _shade_cell(cell, UNICT_LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            _format_run(run, size=9, color=UNICT_BLUE, bold=True)
    for item in items:
        cells = table.add_row().cells
        values = (
            item.criterion_code,
            item.title or item.document_type,
            f"v{item.document_version}",
            _resolution_label(item.resolution_reason),
        )
        for cell, value in zip(cells, values, strict=True):
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_table_widths(table, (0.75, 3.35, 0.75, 1.75))


def _add_annotated_syllabus(
    doc: DocxDocument,
    bundle: EvaluationExportBundle,
) -> None:
    annotations = _annotations_by_field(bundle.evaluation.agent_outputs)
    if not _all_annotations(annotations):
        return
    # Collapse each criterion to a single best anchor so we render only the
    # sections that actually carry a margin note.
    annotations = _filtered_annotation_map(annotations, bundle.syllabus)

    doc.add_page_break()
    doc.add_heading("Syllabus annotato", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "Sono riportate solo le sezioni del syllabus che presentano criticità o "
        "aree migliorabili, con le relative annotazioni a margine. Le sezioni "
        "ritenute adeguate sono omesse per sintesi."
    )
    _format_run(run, size=9, color=MUTED, italic=True)

    added: set[str] = set()
    _add_language_syllabus(doc, bundle.syllabus, "it", annotations, added)
    if bundle.syllabus.has_english:
        _add_language_syllabus(doc, bundle.syllabus, "en", annotations, added)


def _add_language_syllabus(
    doc: DocxDocument,
    syllabus: Any,
    language: str,
    annotations: dict[str, list[Annotation]],
    added: set[str],
) -> None:
    sections = _syllabus_sections(syllabus, language)
    relevant = _relevant_section_keys(sections, annotations, language)
    if not relevant:
        return
    heading_added = False
    for section_key, title, fields in sections:
        if section_key not in relevant:
            continue
        available = [(field, value) for field, value in fields if _has_text(value)]
        if not available:
            continue
        if not heading_added:
            language_name = (
                "Versione italiana" if language == "it" else "English version"
            )
            doc.add_heading(language_name, level=2)
            heading_added = True
        heading = doc.add_paragraph(style="Heading 3")
        heading_run = heading.add_run(title)

        section_annotations = [
            annotation
            for code in SECTION_FALLBACK_CRITERIA.get(section_key, ())
            for annotation in annotations.get("__unanchored__", [])
            if annotation.code == code and annotation.code not in added
        ]
        for field, text in available:
            direct = [
                item for item in annotations.get(field, [])
                if item.code not in added
            ]
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.keep_together = False
            anchored = _add_text_with_comments(doc, paragraph, str(text), direct, added)
            if not anchored and section_annotations:
                for annotation in section_annotations:
                    _add_comment(doc, [heading_run], annotation)
                    added.add(annotation.code)
                section_annotations = []


def _relevant_section_keys(
    sections: list[tuple[str, str, list[tuple[str, Any]]]],
    annotations: dict[str, list[Annotation]],
    language: str,
) -> set[str]:
    """Section keys that will host at least one margin note."""
    note_fields = {field for field in annotations if field != "__unanchored__"}
    keys = {
        section_key
        for section_key, _title, fields in sections
        if any(field in note_fields for field, _ in fields)
    }
    if language == "it":
        unanchored = {a.code for a in annotations.get("__unanchored__", [])}
        keys |= {
            section_key
            for section_key, codes in SECTION_FALLBACK_CRITERIA.items()
            if any(code in unanchored for code in codes)
        }
    return keys


def _filtered_annotation_map(
    annotations: dict[str, list[Annotation]],
    syllabus: Any,
) -> dict[str, list[Annotation]]:
    """Keep one anchor per criterion, grouped by its target field."""
    out: dict[str, list[Annotation]] = {}
    for annotation in _primary_annotations(annotations, syllabus).values():
        key = annotation.source_field or "__unanchored__"
        out.setdefault(key, []).append(annotation)
    return out


def _primary_annotations(
    annotations: dict[str, list[Annotation]],
    syllabus: Any,
) -> dict[str, Annotation]:
    """Pick the single best field anchor for each annotated criterion.

    Prefers a field that is actually rendered, whose evidence text is found
    in it, and (tie-break) the Italian version. Criteria with no usable field
    anchor fall back to the section heading via ``__unanchored__``.
    """
    rendered = _rendered_fields(syllabus)
    chosen: dict[str, Annotation] = {}
    for field, items in annotations.items():
        if field == "__unanchored__" or field not in rendered:
            continue
        for annotation in items:
            if _better_anchor(annotation, chosen.get(annotation.code), syllabus):
                chosen[annotation.code] = annotation
    for annotation in annotations.get("__unanchored__", []):
        chosen.setdefault(annotation.code, annotation)
    for annotation in _all_annotations(annotations):
        if annotation.code not in chosen:
            chosen[annotation.code] = Annotation(
                code=annotation.code,
                score=annotation.score,
                justification=annotation.justification,
                evidence_text=None,
                source_field=None,
            )
    return chosen


def _better_anchor(new: Annotation, old: Annotation | None, syllabus: Any) -> bool:
    if old is None:
        return True
    new_match = _evidence_in_field(new, syllabus)
    old_match = _evidence_in_field(old, syllabus)
    if new_match != old_match:
        return new_match
    new_it = (new.source_field or "").endswith("_it")
    old_it = (old.source_field or "").endswith("_it")
    if new_it != old_it:
        return new_it
    return False


def _evidence_in_field(annotation: Annotation, syllabus: Any) -> bool:
    if not annotation.evidence_text or not annotation.source_field:
        return False
    value = getattr(syllabus, annotation.source_field, None)
    return _has_text(value) and annotation.evidence_text in str(value)


def _rendered_fields(syllabus: Any) -> set[str]:
    fields: set[str] = set()
    for language in ("it", "en"):
        for _key, _title, flds in _syllabus_sections(syllabus, language):
            for field, value in flds:
                if _has_text(value):
                    fields.add(field)
    return fields


def _add_text_with_comments(
    doc: DocxDocument,
    paragraph: Any,
    text: str,
    annotations: list[Annotation],
    added: set[str],
) -> bool:
    pending = [item for item in annotations if item.code not in added]
    if not pending:
        paragraph.add_run(text)
        return False

    cursor = 0
    anchored = False
    for annotation in pending:
        evidence = annotation.evidence_text or ""
        index = text.find(evidence, cursor) if evidence else -1
        if index < 0:
            continue
        if index > cursor:
            paragraph.add_run(text[cursor:index])
        target = paragraph.add_run(text[index:index + len(evidence)])
        _highlight_annotation_run(target, annotation.score)
        _add_comment(doc, [target], annotation)
        added.add(annotation.code)
        anchored = True
        cursor = index + len(evidence)
    if cursor < len(text):
        paragraph.add_run(text[cursor:])
    if not paragraph.runs:
        paragraph.add_run(text)
    leftovers = [item for item in pending if item.code not in added]
    if leftovers:
        target = paragraph.runs[-1]
        for annotation in leftovers:
            _add_comment(doc, [target], annotation)
            added.add(annotation.code)
        anchored = True
    return anchored


def _add_method_note(doc: DocxDocument) -> None:
    doc.add_heading("Note metodologiche", level=1)
    notes = (
        "Il sistema è uno strumento di supporto alla valutazione e "
        "all'autovalutazione del syllabus, non un sostituto del giudizio del "
        "docente o degli organi di Assicurazione della Qualità.",
        "Le evidenze citate e i giudizi prodotti vanno verificati sul syllabus "
        "originale: il sistema lavora su dati estratti tramite scraping e "
        "parsing, e residui di estrazione (refusi apparenti, interruzioni di "
        "parola anomale, marker di template) possono essere artefatti "
        "dell'estrazione e non difetti del documento redatto dal docente.",
    )
    for text in notes:
        p = doc.add_paragraph()
        _format_run(p.add_run(text), size=9, color=MUTED)


def _core_judgments(raw_outputs: Any) -> dict[str, dict[str, Any]]:
    out: dict[str, dict[str, Any]] = {}
    if not isinstance(raw_outputs, dict):
        return out
    for agent in raw_outputs.values():
        if not isinstance(agent, dict):
            continue
        for judgment in agent.get("judgments") or []:
            if isinstance(judgment, dict) and judgment.get("criterion_code"):
                out[str(judgment["criterion_code"])] = judgment
    return out


def _annotations_by_field(raw_outputs: Any) -> dict[str, list[Annotation]]:
    out: dict[str, list[Annotation]] = {}
    for code, judgment in _core_judgments(raw_outputs).items():
        score = _score_value(judgment.get("score"))
        if score not in (0, 1):
            continue
        evidences = judgment.get("evidences") or []
        anchored = False
        if evidences:
            for evidence in evidences:
                if not isinstance(evidence, dict):
                    continue
                field = evidence.get("source_field")
                if not field:
                    continue
                anchored = True
                out.setdefault(str(field), []).append(
                    Annotation(
                        code=code,
                        score=score,
                        justification=str(judgment.get("justification") or ""),
                        evidence_text=str(evidence.get("text") or "") or None,
                        source_field=str(field),
                    )
                )
        if not anchored:
            out.setdefault("__unanchored__", []).append(
                Annotation(
                    code=code,
                    score=score,
                    justification=str(judgment.get("justification") or ""),
                    evidence_text=None,
                    source_field=None,
                )
            )
    return out


def _all_annotations(mapping: dict[str, list[Annotation]]) -> list[Annotation]:
    unique: dict[str, Annotation] = {}
    for items in mapping.values():
        for item in items:
            unique.setdefault(item.code, item)
    return list(unique.values())


def _add_comment(doc: DocxDocument, runs: list[Any], annotation: Annotation) -> None:
    label = "Criticità" if annotation.score == 0 else "Area da migliorare"
    recommendation = CRITERION_RECOMMENDATIONS.get(annotation.code)
    parts = [
        f"{annotation.code} · {CRITERION_NAMES.get(annotation.code, annotation.code)}",
        f"Esito: {label}",
        _humanize(annotation.justification),
    ]
    if recommendation:
        parts.append(f"Indicazione: {recommendation}")
    doc.add_comment(
        runs,
        text="\n\n".join(part for part in parts if part),
        author="Syllabus Quality Assurance",
        initials="SQA",
    )


def _syllabus_sections(
    syllabus: Any,
    language: str,
) -> list[tuple[str, str, list[tuple[str, Any]]]]:
    suffix = f"_{language}"
    learning_fields = _learning_outcome_fields(syllabus, language)
    return [
        (
            "learning_outcomes",
            "Risultati di apprendimento attesi" if language == "it"
            else "Expected learning outcomes",
            learning_fields,
        ),
        (
            "teaching_methods",
            "Modalità di svolgimento dell'insegnamento" if language == "it"
            else "Teaching methods",
            [(f"teaching_methods{suffix}", getattr(syllabus, f"teaching_methods{suffix}"))],
        ),
        (
            "prerequisites",
            "Prerequisiti richiesti" if language == "it" else "Prerequisites",
            [(f"prerequisites{suffix}", getattr(syllabus, f"prerequisites{suffix}"))],
        ),
        (
            "attendance",
            "Frequenza lezioni" if language == "it" else "Attendance",
            [(f"attendance{suffix}", getattr(syllabus, f"attendance{suffix}"))],
        ),
        (
            "course_content",
            "Contenuti del corso" if language == "it" else "Course content",
            [(f"course_content{suffix}", getattr(syllabus, f"course_content{suffix}"))],
        ),
        (
            "references",
            "Testi di riferimento" if language == "it" else "References",
            [(f"references{suffix}", getattr(syllabus, f"references{suffix}"))],
        ),
        (
            "schedule",
            "Programmazione del corso" if language == "it" else "Course schedule",
            [(f"schedule{suffix}", _schedule_text(getattr(syllabus, f"schedule{suffix}")))],
        ),
        (
            "assessment_methods",
            "Modalità di verifica dell'apprendimento" if language == "it"
            else "Assessment methods",
            [(f"assessment_methods{suffix}", getattr(syllabus, f"assessment_methods{suffix}"))],
        ),
        (
            "sample_questions",
            "Esempi di domande e/o esercizi frequenti" if language == "it"
            else "Sample questions",
            [(f"sample_questions{suffix}", getattr(syllabus, f"sample_questions{suffix}"))],
        ),
    ]


def _learning_outcome_fields(
    syllabus: Any,
    language: str,
) -> list[tuple[str, Any]]:
    suffix = f"_{language}"
    descriptors = [
        (
            f"dublin_knowledge{suffix}",
            "Conoscenza e capacità di comprensione"
            if language == "it" else "Knowledge and understanding",
        ),
        (
            f"dublin_applying{suffix}",
            "Capacità di applicare conoscenza e comprensione"
            if language == "it" else "Applying knowledge and understanding",
        ),
        (
            f"dublin_judgement{suffix}",
            "Autonomia di giudizio"
            if language == "it" else "Making judgements",
        ),
        (
            f"dublin_communication{suffix}",
            "Abilità comunicative" if language == "it" else "Communication skills",
        ),
        (
            f"dublin_learning{suffix}",
            "Capacità di apprendimento" if language == "it" else "Learning skills",
        ),
    ]
    populated = [
        (field, f"{label}: {value}")
        for field, label in descriptors
        if _has_text(value := getattr(syllabus, field))
    ]
    if populated:
        return populated
    field = f"learning_outcomes{suffix}"
    return [(field, getattr(syllabus, field))]


def _schedule_text(items: Any) -> str:
    if not isinstance(items, list):
        return ""
    lines: list[str] = []
    for index, item in enumerate(items, 1):
        if not isinstance(item, dict):
            continue
        topic = (
            item.get("argomenti") or item.get("subjects") or item.get("subject")
            or item.get("topics") or item.get("topic") or ""
        )
        references = (
            item.get("riferimenti_testi") or item.get("text_references")
            or item.get("textbook_references") or item.get("references") or ""
        )
        line = f"{item.get('numero') or index}. {topic}"
        if references:
            line += f" — {references}"
        lines.append(line)
    return "\n".join(lines)


_OUTCOME_TOKENS: dict[int | None, tuple[str, RGBColor]] = {
    0: ("Criticità", CRITICAL),
    1: ("Da migliorare", IMPROVABLE),
    2: ("Adeguato", ADEQUATE),
    None: ("Non valutabile", MUTED),
}


def _outcome_tokens(score: int | None) -> tuple[str, RGBColor]:
    return _OUTCOME_TOKENS[score]


def _add_outcome_run(paragraph: Any, score: int | None) -> None:
    label, color = _outcome_tokens(score)
    text = f"  ·  {score}/2 · {label}" if score is not None else f"  ·  {label}"
    run = paragraph.add_run(text)
    _format_run(run, size=9, color=color, bold=True)


def _priority_codes(scores: Any) -> list[tuple[str, int]]:
    if not isinstance(scores, dict):
        return []
    return [
        (code, value)
        for code in CORE_ORDER
        if (value := _score_value(scores.get(code))) in (0, 1)
    ]


def _na_reason_map(items: Any) -> dict[str, str]:
    if not isinstance(items, list):
        return {}
    return {
        str(item.get("criterion_code")): str(item.get("reason") or "")
        for item in items
        if isinstance(item, dict) and item.get("criterion_code")
    }


def _score_value(value: Any) -> int | None:
    return value if value in (0, 1, 2) else None


def _evaluated_count(scores: Any) -> str:
    count = sum(
        _score_value(value) is not None
        for value in scores.values()
    ) if isinstance(scores, dict) else 0
    return f"{count}/9"


def _format_score(value: Any) -> str:
    return f"{float(value):.2f}/2" if value is not None else "—"


def _format_percent(value: Any) -> str:
    return f"{round(float(value) * 100)}%" if value is not None else "—"


def _status_label(status: str) -> str:
    return {
        "completed": "Completata",
        "partial": "Parziale",
        "failed": "Non riuscita",
    }.get(status, status)


def _resolution_label(reason: str) -> str:
    return {
        "explicit_selection": "Selezione esplicita",
        "academic_year_match": "Anno accademico",
        "latest_available_fallback": "Ultima versione disponibile",
    }.get(reason, reason)


def _format_date(value: datetime) -> str:
    return value.strftime("%d/%m/%Y")


def _as_date(value: datetime) -> date:
    return value.date()


def _filename_part(value: str) -> str:
    ascii_value = (
        unicodedata.normalize("NFKD", value)
        .encode("ascii", "ignore")
        .decode("ascii")
    )
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", ascii_value.strip())
    return cleaned.strip("_") or "evaluation"


# Database field names the LLM agents sometimes cite verbatim in their
# justifications. They must never reach a human reader as raw identifiers.
_FIELD_LABELS: dict[str, str] = {
    "course_name_it": "il titolo del corso in italiano",
    "course_name_en": "il titolo del corso in inglese",
    "learning_outcomes_it": "i risultati di apprendimento (versione italiana)",
    "learning_outcomes_en": "i risultati di apprendimento (versione inglese)",
    "dublin_knowledge_it": "il descrittore «Conoscenza e comprensione» (IT)",
    "dublin_knowledge_en": "il descrittore «Conoscenza e comprensione» (EN)",
    "dublin_applying_it": "il descrittore «Capacità di applicare conoscenza» (IT)",
    "dublin_applying_en": "il descrittore «Capacità di applicare conoscenza» (EN)",
    "dublin_judgement_it": "il descrittore «Autonomia di giudizio» (IT)",
    "dublin_judgement_en": "il descrittore «Autonomia di giudizio» (EN)",
    "dublin_communication_it": "il descrittore «Abilità comunicative» (IT)",
    "dublin_communication_en": "il descrittore «Abilità comunicative» (EN)",
    "dublin_learning_it": "il descrittore «Capacità di apprendimento» (IT)",
    "dublin_learning_en": "il descrittore «Capacità di apprendimento» (EN)",
}


_FIELD_TOKEN = r"`?[A-Za-z][A-Za-z0-9]*(?:_[A-Za-z0-9]+)+`?"


def _humanize(text: str) -> str:
    """Strip leaked DB field tokens from agent prose.

    The agents already gloss the field in plain language next to the token
    (``il titolo del corso in inglese (`course_name_en`)``), so the token is
    redundant noise: drop the parenthetical or ``nel campo …`` reference, and
    relabel only a bare token that has no human gloss.
    """
    if not text:
        return text

    # Parenthetical aside, e.g. "(`course_name_en`)".
    text = re.sub(rf"\s*\(\s*{_FIELD_TOKEN}\s*\)", "", text)
    # Inline reference, e.g. "nel campo `dublin_communication_it`".
    text = re.sub(
        rf"\s*(?:nel|nei|del|dei|sul|sui|al|ai|il|i)?\s*camp[oi]\s+"
        rf"(?:dedicat[oi]\s+)?{_FIELD_TOKEN}",
        "",
        text,
    )

    def _sub(match: re.Match[str]) -> str:
        key = match.group(1)
        return _FIELD_LABELS.get(key, key.replace("_", " "))

    # Any surviving token (no nearby gloss) becomes a readable label.
    text = re.sub(r"`([A-Za-z_][A-Za-z0-9_]*)`", _sub, text)
    for field, label in _FIELD_LABELS.items():
        text = text.replace(field, label)

    text = re.sub(r"\s+([,.;:])", r"\1", text)
    return re.sub(r"\s{2,}", " ", text).strip()


def _humanize_na_reason(reason: str | None) -> str:
    """Turn an internal NA message into a teacher-facing sentence."""
    text = (reason or "").strip()
    low = text.lower()
    if not text:
        return "Motivazione non disponibile."
    if "no indexed document" in low:
        return (
            "Per questo Corso di Studio non è disponibile il documento di "
            "riferimento richiesto dal criterio."
        )
    if "no judgment produced" in low or "not applicable" in low:
        return "Il sistema non ha prodotto un giudizio per questo criterio."
    text = re.sub(r"\s*on cdl_id=\d+", "", text)
    return _humanize(text)


def _select_evidences(evidences: Any, seen: set[str]) -> list[str]:
    """Up to three clean, de-duplicated evidence quotes for a criterion."""
    rendered: list[str] = []
    if not isinstance(evidences, list):
        return rendered
    for evidence in evidences:
        if not isinstance(evidence, dict):
            continue
        text = _evidence_text(str(evidence.get("text", "")))
        if not text:
            continue
        key = " ".join(text.lower().split())
        if key in seen:
            continue
        seen.add(key)
        rendered.append(_truncate(text, 360))
        if len(rendered) >= 3:
            break
    return rendered


def _evidence_text(raw: str) -> str:
    """Render a raw evidence quote, unpacking schedule-JSON dumps."""
    text = (raw or "").strip()
    looks_like_schedule = text.startswith("[") and any(
        token in text for token in ('"argomenti"', '"numero"', '"subjects"', '"subject"')
    )
    if looks_like_schedule:
        try:
            data = json.loads(text)
        except (ValueError, TypeError):
            data = None
        if isinstance(data, list):
            readable = _schedule_text(data)
            if readable:
                text = readable.replace("\n", " · ")
    return text


def _truncate(text: str, limit: int) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + "…"


def _first_sentence(text: str, limit: int = 240) -> str:
    """First sentence of a justification, for the adequate-criteria recap."""
    text = _humanize(text or "").strip()
    if not text:
        return ""
    match = re.search(r"[.!?]\s", text)
    sentence = text[: match.start() + 1] if match else text
    return _truncate(sentence, limit)


def _has_text(value: Any) -> bool:
    return isinstance(value, str) and bool(value.strip())


def _format_run(
    run: Any,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_widths(table: Any, widths_inches: tuple[float, ...]) -> None:
    """Apply fixed Word table geometry; all widths sum to 6.6 inches."""
    table.autofit = False
    widths_dxa = [round(width * 1440) for width in widths_inches]
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_dxa, strict=True):
        col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def _highlight_annotation_run(run: Any, score: int) -> None:
    run.font.highlight_color = (
        WD_COLOR_INDEX.YELLOW if score == 1 else WD_COLOR_INDEX.RED
    )
