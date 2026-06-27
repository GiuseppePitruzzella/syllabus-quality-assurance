"""Text extractors for the local-document registry (Phase 8.B.1).

Dispatch by extension and return a deterministic, normalized
plain-text version of the document. All failure modes — unsupported
extension, unreadable file, parser-level exception, scanned PDF whose
optional OCR fallback fails, empty / whitespace-only output, path
escaping the storage root — surface as a single `ExtractionError`. Library
exceptions are wrapped via `raise ... from e` so the original cause
is preserved internally without leaking the stack trace at the API
layer.

The module is intentionally free of any FastAPI / SQLAlchemy
dependency: Phase 8.B.2 will call it from the async indexing job,
Phase 9 will not call it at all (extraction lives in the registry
layer, not in the evaluation graph).
"""
from __future__ import annotations

from pathlib import Path
from typing import Callable

import pdfplumber
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


class ExtractionError(Exception):
    """Raised when text extraction fails or yields no usable text.

    The original library exception is preserved as ``__cause__``
    (via ``raise ... from e``). Callers may log it in full, but
    should not bubble it up unfiltered to API responses.
    """


# ---------------------------------------------------------------------------
# Path safety
# ---------------------------------------------------------------------------


def resolve_local_document_path(
    rel_path: str | Path, base_dir: str | Path,
) -> Path:
    """Resolve a relative document path under the storage root.

    The relative path comes from the DB and is normally generated
    server-side (`{cdl_id}/{id}_v{version}{ext}`), so it should
    never contain `..` segments. We resolve it anyway and assert
    containment under ``base_dir`` so a hypothetical corrupted /
    malicious row can't be turned into a filesystem-escape primitive
    by a future caller.

    Raises:
        ExtractionError: if the resolved path escapes ``base_dir``.
    """
    base = Path(base_dir).resolve()
    candidate = (base / rel_path).resolve()
    try:
        candidate.relative_to(base)
    except ValueError as e:
        raise ExtractionError(
            f"Resolved path {candidate} escapes storage root {base}",
        ) from e
    return candidate


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------


def extract_text(
    abs_path: Path,
    extension: str,
    *,
    pdf_ocr: Callable[[Path], str] | None = None,
) -> str:
    """Extract normalized plain text from a local document.

    ``abs_path`` must already be the resolved absolute path
    (see :func:`resolve_local_document_path`). ``extension`` is the
    file extension as recorded on the registry row (e.g. ``.pdf``);
    ``.htm`` is normalised to ``.html`` so callers can pass either.

    ``pdf_ocr`` is invoked only when a syntactically valid PDF yields
    no text through the standard extractor. It stays unused for
    ordinary text PDFs.

    Returns the normalized text.

    Raises:
        ExtractionError: for any failure (unsupported extension,
            unreadable file, library exception, empty result).
    """
    ext = extension.lower()
    if ext == ".htm":
        ext = ".html"
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise ExtractionError(f"Unsupported extension: {extension!r}")
    if not abs_path.is_file():
        raise ExtractionError(f"File not found: {abs_path}")
    raw = extractor(abs_path)
    try:
        return _normalize_or_raise(raw, extension=ext)
    except ExtractionError:
        if ext != ".pdf" or pdf_ocr is None:
            raise

    try:
        ocr_raw = pdf_ocr(abs_path)
    except Exception as exc:
        raise ExtractionError(f"PDF OCR fallback failed: {exc}") from exc
    try:
        return _normalize_or_raise(ocr_raw, extension=ext)
    except ExtractionError as exc:
        raise ExtractionError(
            "PDF OCR fallback completed but produced no usable text."
        ) from exc


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF, page by page, joined by ``\\n\\n``.

    A PDF whose pages all return empty / None (e.g. a scan without
    OCR) leaves the raw text whitespace-only; downstream
    :func:`_normalize_or_raise` turns that into an explicit
    ``ExtractionError`` with a scan-friendly message rather than
    silently producing an empty string.
    """
    try:
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    except Exception as e:
        raise ExtractionError(f"Failed to read PDF: {e}") from e
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """Extract paragraphs + linearised tables from a DOCX.

    Tables become one row per output line, cells joined by ` | `.
    Empty paragraphs / cells are skipped so the normalised output
    doesn't carry trivia from the source layout.
    """
    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        raise ExtractionError(f"Failed to read DOCX: {e}") from e

    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_html(path: Path) -> str:
    """Extract text from HTML/HTM, dropping non-content elements.

    `script`, `style` and `noscript` are decomposed before
    text extraction so a CdL page with inline JS or CSS doesn't
    pollute the chunked output. BeautifulSoup is configured with
    the lxml parser already available repo-wide.
    """
    try:
        raw = path.read_bytes()
        soup = BeautifulSoup(raw, "html.parser")
    except Exception as e:
        raise ExtractionError(f"Failed to read HTML: {e}") from e
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n")


def _extract_md(path: Path) -> str:
    """Read Markdown as plain text, UTF-8 with latin-1 fallback.

    No Markdown parsing: per Phase 8 brief, the raw text is the
    document. Markdown is a transport format here, not a structured
    schema we care about.
    """
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="latin-1")
        except Exception as e:
            raise ExtractionError(f"Failed to read Markdown: {e}") from e
    except OSError as e:
        raise ExtractionError(f"Failed to read Markdown: {e}") from e


_EXTRACTORS: dict[str, Callable[[Path], str]] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".html": _extract_html,
    ".md": _extract_md,
}


# ---------------------------------------------------------------------------
# Normalisation
# ---------------------------------------------------------------------------


def _normalize_or_raise(raw: str, extension: str) -> str:
    """Deterministic normalisation; raise on empty result.

    Steps:
      - normalise line endings to ``\\n``
      - trim trailing whitespace per line
      - collapse runs of 2+ blank lines to a single blank line
      - strip leading/trailing whitespace
      - raise ExtractionError if nothing remains

    The PDF branch is special-cased in the error message because
    "empty after extraction" usually means the PDF is a scan
    without OCR and the user needs a specific hint.
    """
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    out: list[str] = []
    blank_run = False
    for line in lines:
        if line.strip() == "":
            if not blank_run:
                out.append("")
                blank_run = True
        else:
            out.append(line)
            blank_run = False
    normalized = "\n".join(out).strip()
    if not normalized:
        if extension == ".pdf":
            raise ExtractionError(
                "PDF contains no extractable text. "
                "If this is a scan, OCR is required before upload.",
            )
        raise ExtractionError(
            "Extracted text is empty after normalisation.",
        )
    return normalized
