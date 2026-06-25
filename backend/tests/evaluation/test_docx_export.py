from __future__ import annotations

from datetime import datetime, timedelta, timezone
from io import BytesIO
from types import SimpleNamespace
from zipfile import ZipFile

import pytest
from docx import Document

from app.evaluation.docx_export import (
    Annotation,
    EvaluationExportBundle,
    EvaluationExportNotFoundError,
    _add_comment,
    _add_run,
    _evidence_text,
    _humanize,
    _humanize_na_reason,
    _primary_annotations,
    _select_evidences,
    build_evaluation_docx,
    export_filename,
    load_cdl_export_bundles,
    load_evaluation_export_bundle,
)
from app.models import CorsoDiLaurea, Department, EvaluationResult, Syllabus


def test_build_docx_contains_evaluation_sections_and_real_comment():
    bundle = _synthetic_bundle()

    payload = build_evaluation_docx(bundle)

    doc = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Valutazione della qualità del syllabus" in text
    assert "Sintesi della valutazione" in text
    assert "Revisione dei criteri C1-C9" in text
    assert "Criteri estesi" in text
    assert "Syllabus annotato" in text
    assert "Note metodologiche" in text
    # The legacy synthesizer report must no longer be duplicated into the DOCX.
    assert "Report di valutazione" not in text
    assert "Dettagli tecnici" not in text
    # Scorecard: a scannable C1-C9 overview table with score + outcome.
    assert "Punteggio" in table_text
    assert "Adeguato" in table_text
    assert "Da migliorare" in table_text
    with ZipFile(BytesIO(payload)) as archive:
        assert "word/comments.xml" in archive.namelist()
        comments = archive.read("word/comments.xml").decode()
        document = archive.read("word/document.xml").decode()
    assert "C5 · Chiarezza dei prerequisiti" in comments
    assert "Indicazione:" in comments
    assert "commentRangeStart" in document
    assert "commentRangeEnd" in document
    assert "commentReference" in document


def test_humanize_replaces_leaked_db_field_tokens():
    out = _humanize(
        "Il `course_name_en` è assente e dublin_communication_it è troncato."
    )
    assert "course_name_en" not in out
    assert "dublin_communication_it" not in out
    assert "titolo del corso in inglese" in out
    assert "Abilità comunicative" in out


def test_humanize_drops_redundant_field_parenthetical():
    out = _humanize(
        "il titolo del corso in inglese (`course_name_en`) risulta assente"
    )
    assert out == "il titolo del corso in inglese risulta assente"


def test_humanize_drops_inline_campo_field_reference():
    out = _humanize(
        "il descrittore «Abilità comunicative» nel campo "
        "`dublin_communication_it` risulta troncato"
    )
    assert "dublin_communication_it" not in out
    assert "nel campo" not in out
    assert out == "il descrittore «Abilità comunicative» risulta troncato"


def test_humanize_na_reason_hides_internal_identifiers():
    resolver = _humanize_na_reason("no indexed document enabled for E1 on cdl_id=3")
    assert "cdl_id" not in resolver
    assert "indexed" not in resolver.lower()

    handler = _humanize_na_reason("no judgment produced")
    assert "no judgment" not in handler.lower()

    assert _humanize_na_reason(None) == "Motivazione non disponibile."


def test_comment_humanizes_internal_field_names():
    doc = Document()
    run = doc.add_paragraph().add_run("Titolo")
    _add_comment(
        doc,
        [run],
        Annotation(
            code="C2",
            score=1,
            justification="Il `course_name_en` risulta assente.",
            evidence_text="Titolo",
            source_field="course_name_it",
        ),
    )
    buffer = BytesIO()
    doc.save(buffer)
    with ZipFile(BytesIO(buffer.getvalue())) as archive:
        comments = archive.read("word/comments.xml").decode()
    assert "course_name_en" not in comments
    assert "titolo del corso in inglese" in comments


def test_primary_annotations_falls_back_when_source_field_is_not_rendered():
    syllabus = SimpleNamespace(
        **{
            field: None
            for field in (
                "learning_outcomes_it",
                "learning_outcomes_en",
                "dublin_knowledge_it",
                "dublin_knowledge_en",
                "dublin_applying_it",
                "dublin_applying_en",
                "dublin_judgement_it",
                "dublin_judgement_en",
                "dublin_communication_it",
                "dublin_communication_en",
                "dublin_learning_it",
                "dublin_learning_en",
                "teaching_methods_it",
                "teaching_methods_en",
                "prerequisites_it",
                "prerequisites_en",
                "attendance_it",
                "attendance_en",
                "course_content_it",
                "course_content_en",
                "references_it",
                "references_en",
                "schedule_it",
                "schedule_en",
                "assessment_methods_it",
                "assessment_methods_en",
                "sample_questions_it",
                "sample_questions_en",
            )
        }
    )
    annotation = Annotation(
        code="C9",
        score=1,
        justification="Motivazione leggibile.",
        evidence_text="Testo",
        source_field="syllabus.references_it",
    )
    selected = _primary_annotations(
        {"syllabus.references_it": [annotation]},
        syllabus,
    )
    assert selected["C9"].source_field is None
    assert selected["C9"].evidence_text is None


def test_add_run_converts_newlines_to_line_breaks():
    doc = Document()
    run = _add_run(doc.add_paragraph(), "1. Primo\n2. Secondo\n3. Terzo")
    xml = run._element.xml
    assert xml.count("<w:br/>") == 2
    assert "Primo" in xml and "Terzo" in xml


def test_select_evidences_drops_null_placeholders():
    out = _select_evidences(
        [
            {"text": "null"},
            {"text": "None"},
            {"text": ""},
            {"text": "  N/A "},
            {"text": "Contenuto reale del syllabus"},
        ],
        set(),
    )
    assert out == ["Contenuto reale del syllabus"]


def test_extended_section_hidden_when_no_criterion_evaluated():
    bundle = _synthetic_bundle()
    bundle.evaluation.extended_criteria_result = {
        "criterion_scores": {code: None for code in ("E1", "E2", "E3", "E4", "E5")},
        "na_criteria": [
            {"criterion_code": "E1", "reason": "no indexed document enabled for E1"},
        ],
        "agent_output": {"judgments": []},
    }

    doc = Document(BytesIO(build_evaluation_docx(bundle)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Criteri estesi" not in text
    assert "Non valutabile" not in text


def test_evidence_text_unpacks_schedule_json_dump():
    raw = (
        '[{"numero": "1", "argomenti": "Intro", "riferimenti_testi": "Cap. 1"}, '
        '{"numero": "2", "argomenti": "Avanzato", "riferimenti_testi": "Cap. 2"}]'
    )
    out = _evidence_text(raw)
    assert "{" not in out
    assert "argomenti" not in out
    assert "Intro" in out and "Avanzato" in out


def test_export_filename_is_stable_and_safe():
    bundle = _synthetic_bundle()

    assert export_filename(bundle) == (
        "LM_18__Deep_Learning__2026-06-25__evaluation.docx"
    )


def test_load_bundle_rejects_failed_evaluation(db_session):
    _seed_catalog(db_session)
    _seed_evaluation(db_session, uuid="failed", status="failed", offset=0)

    with pytest.raises(EvaluationExportNotFoundError):
        load_evaluation_export_bundle(db_session, "failed")


def test_cdl_export_uses_latest_non_failed_run_per_syllabus(db_session):
    _seed_catalog(db_session)
    _seed_evaluation(db_session, uuid="old", status="completed", offset=-3)
    _seed_evaluation(db_session, uuid="latest", status="partial", offset=-1)
    _seed_evaluation(db_session, uuid="failed-newer", status="failed", offset=0)

    cdl, bundles = load_cdl_export_bundles(db_session, 1)

    assert cdl.code == "LM-18"
    assert [bundle.evaluation.evaluation_uuid for bundle in bundles] == ["latest"]


def _synthetic_bundle() -> EvaluationExportBundle:
    now = datetime(2026, 6, 25, 10, 0, tzinfo=timezone.utc)
    scores = {f"C{i}": 2 for i in range(1, 10)}
    scores["C5"] = 1
    evaluation = SimpleNamespace(
        course_name_snapshot="Deep Learning",
        status="completed",
        started_at=now,
        core_score=1.89,
        coverage=1.0,
        criterion_scores=scores,
        na_criteria=[],
        agent_outputs={
            "A1": {
                "judgments": [
                    {
                        "criterion_code": "C5",
                        "score": 1,
                        "is_na": False,
                        "justification": "I prerequisiti sono presenti ma migliorabili.",
                        "evidences": [
                            {
                                "text": "Solide basi di Machine Learning",
                                "source_field": "prerequisites_it",
                            }
                        ],
                        "confidence": "high",
                    }
                ]
            }
        },
        extended_criteria_result={
            "criterion_scores": {
                "E1": None, "E2": None, "E3": None, "E4": 2, "E5": 2,
            },
            "na_criteria": [
                {"criterion_code": "E1", "reason": "Documento non disponibile"},
                {"criterion_code": "E2", "reason": "Documento non disponibile"},
                {"criterion_code": "E3", "reason": "Documento non disponibile"},
            ],
            "agent_output": {
                "judgments": [
                    {
                        "criterion_code": "E4",
                        "score": 2,
                        "justification": "Le versioni sono coerenti.",
                        "evidences": [],
                    },
                    {
                        "criterion_code": "E5",
                        "score": 2,
                        "justification": "Gli usi locali sono rispettati.",
                        "evidences": [],
                    },
                ]
            },
        },
        final_report="# Sintesi\n\n- Documento complessivamente adeguato.",
    )
    syllabus_fields = {
        "course_name": "Deep Learning",
        "teacher": "Antonino Furnari",
        "academic_year": "2025/2026",
        "has_english": True,
        "learning_outcomes_it": "Lo studente acquisisce competenze avanzate.",
        "dublin_knowledge_it": "",
        "dublin_applying_it": "",
        "dublin_judgement_it": "",
        "dublin_communication_it": "",
        "dublin_learning_it": "",
        "teaching_methods_it": "Lezioni e laboratorio.",
        "prerequisites_it": "Solide basi di Machine Learning e Python.",
        "attendance_it": "Consigliata.",
        "course_content_it": "Reti neurali e modelli generativi.",
        "references_it": "Goodfellow et al., Deep Learning.",
        "schedule_it": [{"numero": 1, "argomenti": "Introduzione"}],
        "assessment_methods_it": "Prova scritta e progetto.",
        "sample_questions_it": "Descrivere la backpropagation.",
        "learning_outcomes_en": "The student acquires advanced skills.",
        "dublin_knowledge_en": None,
        "dublin_applying_en": None,
        "dublin_judgement_en": None,
        "dublin_communication_en": None,
        "dublin_learning_en": None,
        "teaching_methods_en": "Lectures and laboratory.",
        "prerequisites_en": "Machine Learning and Python.",
        "attendance_en": "Recommended.",
        "course_content_en": "Neural networks and generative models.",
        "references_en": "Goodfellow et al., Deep Learning.",
        "schedule_en": [{"numero": 1, "subjects": "Introduction"}],
        "assessment_methods_en": "Written exam and project.",
        "sample_questions_en": "Describe backpropagation.",
    }
    return EvaluationExportBundle(
        evaluation=evaluation,
        syllabus=SimpleNamespace(**syllabus_fields),
        cdl=SimpleNamespace(code="LM-18", name="Informatica"),
        department=SimpleNamespace(name="Matematica e Informatica"),
    )


def _seed_catalog(session) -> None:
    now = datetime(2026, 6, 25, tzinfo=timezone.utc)
    session.add(
        Department(
            id=1,
            name="DMI",
            area="Scientifica",
            website_url="https://example.test",
            email="dmi@example.test",
            phone="0",
            director="Direttore",
            scraped_at=now,
        )
    )
    session.add(
        CorsoDiLaurea(
            id=1,
            department_id=1,
            name="Informatica",
            code="LM-18",
            type="Magistrale",
            academic_year="2025/2026",
            url="https://example.test/lm18",
            scraped_at=now,
        )
    )
    fields = {
        "id": 1,
        "cdl_id": 1,
        "seuid": "SEUID-1",
        "course_code": "CODE-1",
        "course_name": "Deep Learning",
        "teacher": "Docente",
        "academic_year": "2025/2026",
        "year_of_study": "1",
        "url_it": "https://example.test/it",
        "url_en": "https://example.test/en",
        "has_english": False,
        "scraped_at": now,
    }
    for field in (
        "learning_outcomes_it", "dublin_knowledge_it", "dublin_applying_it",
        "dublin_judgement_it", "dublin_communication_it", "dublin_learning_it",
        "teaching_methods_it", "prerequisites_it", "attendance_it",
        "course_content_it", "references_it", "assessment_methods_it",
        "sample_questions_it",
    ):
        fields[field] = f"Testo {field}"
    session.add(Syllabus(**fields))
    session.commit()


def _seed_evaluation(session, *, uuid: str, status: str, offset: int) -> None:
    started = datetime(2026, 6, 25, tzinfo=timezone.utc) + timedelta(hours=offset)
    session.add(
        EvaluationResult(
            evaluation_uuid=uuid,
            syllabus_id=1,
            syllabus_seuid_snapshot="SEUID-1",
            course_name_snapshot="Deep Learning",
            status=status,
            started_at=started,
            finished_at=started + timedelta(seconds=5),
            duration_ms=5000,
            llm_model="gemini-2.5-flash",
            embedding_model="gemini-embedding-001",
            embedding_dim=3072,
            llm_temperature=0.1,
            llm_max_output_tokens=8192,
            rag_top_k=5,
            rag_final_k=3,
            rag_similarity_threshold=0.6,
            gcp_project_id="test",
            gcp_location="europe-west8",
            prompt_versions={},
            core_score=1.5,
            coverage=1.0,
            criterion_scores={f"C{i}": 2 for i in range(1, 10)},
            na_criteria=[],
            agent_outputs={},
            agent_errors={},
            retrieved_chunks={},
            final_report="Report",
        )
    )
    session.commit()
