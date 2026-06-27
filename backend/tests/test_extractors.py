"""Tests for the local-document extractors (Phase 8.B.1).

Covers:
  - happy paths for PDF, DOCX, HTML/HTM, MD
  - DOCX with a table
  - HTML strips script / style / noscript
  - empty / whitespace-only output → ExtractionError
  - scanned PDF → ExtractionError without OCR, extracted text with fallback
  - corrupted PDF / DOCX → ExtractionError
  - MD UTF-8 happy path + latin-1 fallback
  - unsupported extension → ExtractionError
  - path escape → ExtractionError

Test fixtures are generated on the fly so the suite stays self-
contained and no real syllabus binary lives under `tests/`.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

from app.local_documents.extractors import (
    ExtractionError,
    extract_text,
    resolve_local_document_path,
)


# ---------------------------------------------------------------------------
# Helpers — synthesise tiny test files on disk under tmp_path
# ---------------------------------------------------------------------------


def _make_pdf(path: Path, lines: list[str]) -> None:
    """Write a tiny single-page PDF with one text line per entry."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    y = 800
    for line in lines:
        c.drawString(72, y, line)
        y -= 18
    c.save()
    path.write_bytes(buf.getvalue())


def _make_blank_pdf(path: Path) -> None:
    """A 1-page PDF with no text — proxy for a scan-without-OCR."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.showPage()  # nothing drawn
    c.save()
    path.write_bytes(buf.getvalue())


def _make_docx(
    path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None,
) -> None:
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for r, row in enumerate(table_rows):
            for col, value in enumerate(row):
                table.cell(r, col).text = value
    doc.save(str(path))


# ---------------------------------------------------------------------------
# Happy paths
# ---------------------------------------------------------------------------


def test_extract_pdf_returns_normalized_text(tmp_path):
    path = tmp_path / "doc.pdf"
    _make_pdf(path, ["Linee guida LM-18", "Prerequisiti culturali"])
    text = extract_text(path, ".pdf")
    assert "Linee guida LM-18" in text
    assert "Prerequisiti culturali" in text


def test_extract_docx_paragraphs(tmp_path):
    path = tmp_path / "doc.docx"
    _make_docx(path, ["Intestazione", "Corpo del testo"])
    text = extract_text(path, ".docx")
    assert text == "Intestazione\n\nCorpo del testo"


def test_extract_docx_with_table_linearised(tmp_path):
    path = tmp_path / "doc.docx"
    _make_docx(
        path,
        ["Premessa"],
        table_rows=[
            ["Criterio", "Peso", "Agente"],
            ["C1", "1.0", "A1"],
            ["C9", "1.0", "A4"],
        ],
    )
    text = extract_text(path, ".docx")
    assert "Premessa" in text
    assert "Criterio | Peso | Agente" in text
    assert "C1 | 1.0 | A1" in text
    assert "C9 | 1.0 | A4" in text


def test_extract_html_drops_script_style_noscript(tmp_path):
    path = tmp_path / "doc.html"
    path.write_text(
        """<html><head>
        <style>body { color: red }</style>
        <script>alert('xss')</script>
        <noscript>js disabled</noscript>
        </head><body><p>Linee guida</p><p>Indicazioni redazionali</p></body></html>""",
        encoding="utf-8",
    )
    text = extract_text(path, ".html")
    assert "Linee guida" in text
    assert "Indicazioni redazionali" in text
    # Decomposed tags must not appear in the output.
    assert "alert" not in text
    assert "color: red" not in text
    assert "js disabled" not in text


def test_extract_htm_normalised_to_html(tmp_path):
    path = tmp_path / "doc.htm"
    path.write_text("<html><body><p>OK</p></body></html>", encoding="utf-8")
    text = extract_text(path, ".htm")
    assert text == "OK"


def test_extract_md_utf8(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("# Linee guida\n\nCorpo italiano à è ò.", encoding="utf-8")
    text = extract_text(path, ".md")
    assert "Linee guida" in text
    assert "à è ò" in text


def test_extract_md_latin1_fallback(tmp_path):
    path = tmp_path / "doc.md"
    # Bytes that are valid latin-1 but invalid UTF-8 (à è ò in latin-1).
    path.write_bytes(b"# Linee guida\n\nCorpo: \xe0 \xe8 \xf2.")
    text = extract_text(path, ".md")
    assert "Linee guida" in text
    assert "à è ò" in text


# ---------------------------------------------------------------------------
# Normalisation
# ---------------------------------------------------------------------------


def test_normalisation_collapses_blank_runs(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("A\n\n\n\nB\n   \n\nC", encoding="utf-8")
    text = extract_text(path, ".md")
    assert text == "A\n\nB\n\nC"


def test_normalisation_strips_trailing_whitespace_per_line(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("A   \nB\t\t\n", encoding="utf-8")
    text = extract_text(path, ".md")
    assert text == "A\nB"


# ---------------------------------------------------------------------------
# Error paths
# ---------------------------------------------------------------------------


def test_empty_md_raises(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("", encoding="utf-8")
    with pytest.raises(ExtractionError):
        extract_text(path, ".md")


def test_whitespace_only_md_raises(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("   \n\n  \t\n", encoding="utf-8")
    with pytest.raises(ExtractionError):
        extract_text(path, ".md")


def test_scanned_pdf_without_text_raises(tmp_path):
    path = tmp_path / "scan.pdf"
    _make_blank_pdf(path)
    with pytest.raises(ExtractionError) as excinfo:
        extract_text(path, ".pdf")
    # The message should mention OCR / scan so the API can surface a
    # specific hint to the user.
    assert "OCR" in str(excinfo.value) or "scan" in str(excinfo.value).lower()


def test_scanned_pdf_uses_injected_ocr_fallback(tmp_path):
    path = tmp_path / "scan.pdf"
    _make_blank_pdf(path)
    calls: list[Path] = []

    def fake_ocr(received: Path) -> str:
        calls.append(received)
        return "Risultato | Attività formative\nConoscenza | INF/01"

    text = extract_text(path, ".pdf", pdf_ocr=fake_ocr)

    assert calls == [path]
    assert "Risultato | Attività formative" in text


def test_text_pdf_does_not_call_ocr_fallback(tmp_path):
    path = tmp_path / "text.pdf"
    _make_pdf(path, ["Matrice di Tuning LM-13"])

    def unexpected_ocr(_path: Path) -> str:
        raise AssertionError("OCR must not run for a text PDF")

    assert "Matrice di Tuning" in extract_text(
        path,
        ".pdf",
        pdf_ocr=unexpected_ocr,
    )


def test_scanned_pdf_wraps_ocr_failure(tmp_path):
    path = tmp_path / "scan.pdf"
    _make_blank_pdf(path)

    def failing_ocr(_path: Path) -> str:
        raise RuntimeError("vertex unavailable")

    with pytest.raises(ExtractionError, match="OCR fallback failed"):
        extract_text(path, ".pdf", pdf_ocr=failing_ocr)


def test_corrupted_pdf_raises(tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4 garbage bytes that are not a valid PDF")
    with pytest.raises(ExtractionError):
        extract_text(path, ".pdf")


def test_corrupted_docx_raises(tmp_path):
    path = tmp_path / "broken.docx"
    path.write_bytes(b"PK garbage - not a real docx archive")
    with pytest.raises(ExtractionError):
        extract_text(path, ".docx")


def test_unsupported_extension_raises(tmp_path):
    path = tmp_path / "doc.txt"
    path.write_text("plain text", encoding="utf-8")
    with pytest.raises(ExtractionError):
        extract_text(path, ".txt")


def test_missing_file_raises(tmp_path):
    path = tmp_path / "absent.md"
    with pytest.raises(ExtractionError):
        extract_text(path, ".md")


# ---------------------------------------------------------------------------
# Path safety
# ---------------------------------------------------------------------------


def test_resolve_safe_path_happy(tmp_path):
    (tmp_path / "3").mkdir()
    (tmp_path / "3" / "1_v1.md").write_text("ok", encoding="utf-8")
    resolved = resolve_local_document_path("3/1_v1.md", tmp_path)
    assert resolved.is_file()


def test_resolve_path_rejects_escape_via_dotdot(tmp_path):
    (tmp_path / "store").mkdir()
    base = tmp_path / "store"
    with pytest.raises(ExtractionError):
        resolve_local_document_path("../escape.md", base)


def test_resolve_path_rejects_absolute_outside_root(tmp_path):
    (tmp_path / "store").mkdir()
    base = tmp_path / "store"
    outside = tmp_path / "outside.md"
    outside.write_text("x", encoding="utf-8")
    with pytest.raises(ExtractionError):
        resolve_local_document_path(str(outside), base)
